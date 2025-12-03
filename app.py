import json
import io
from flask import Flask, request, jsonify, render_template, send_from_directory
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import uuid
from googletrans import Translator
from src.translator import TranslationService
from PIL import Image
from docx.shared import Inches

app = Flask(__name__)
translation_service = TranslationService()

translator = Translator()

# Ensure static folder exists to save files
os.makedirs("static", exist_ok=True)

@app.route('/')
def index():
    languages = load_languages()
    languages.sort(key=lambda lang: lang['name'])
    return render_template('index.html', languages=languages)

def load_languages():
    with open('data/languages.json') as f:
        return json.load(f)

@app.route('/detect-language', methods=['POST'])
def detect_language():
    data = request.get_json()
    text = data.get('text')

    detected_language = translation_service.detect_language(text)
    
    return jsonify({
        'language_code': detected_language['code'], 
        'language_name': detected_language['name']  
    })

@app.route('/translate-text', methods=['POST'])
def translate_text():
    data = request.get_json()
    print("Received data for translation:", data)
    source_language = data.get('source-language')
    target_language = data.get('target-language')
    text = data.get('text')

    translated_text = translate(text, src=source_language, dest=target_language)
    print("Translated text:", translated_text) 

    return jsonify({'translated_text': translated_text})

def translate(text, src='auto', dest='en'):
    translation = translator.translate(text, src=src, dest=dest)
    return translation.text

@app.route('/translate-document', methods=['POST'])
def translate_doc():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    uploaded_file = request.files['file']
    if not uploaded_file.filename.endswith('.docx'):
        return jsonify({'error': 'Only Word (.docx) files are allowed'}), 400

    doc = Document(io.BytesIO(uploaded_file.read()))
    translated_doc = Document()
    language_code = request.form.get('language')

    # Process paragraphs and tables
    for element in doc.element.body:
        if element.tag.endswith('tbl'):
            original_table = doc.tables[0]
            translated_table = translated_doc.add_table(rows=0, cols=len(original_table.columns))
            copy_table_style(original_table, translated_table)

            for row in original_table.rows:
                translated_row = translated_table.add_row()
                for cell_idx, cell in enumerate(row.cells):
                    translated_text = translate_text_logic(cell.text, language_code)
                    translated_cell = translated_row.cells[cell_idx]
                    translated_cell.text = translated_text
                    if cell.paragraphs and cell.paragraphs[0].alignment:
                        translated_cell.paragraphs[0].alignment = cell.paragraphs[0].alignment

                    # Check for images inside table cells
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if 'image' in run._r.xml:
                                copy_images(doc, translated_doc, translated_cell)

        elif element.tag.endswith('p'):
            paragraph = next(p for p in doc.paragraphs if p._p == element)
            combined_text = ''.join(run.text for run in paragraph.runs)

            # Translate paragraph text
            translated_text = translate_text_logic(combined_text, language_code)

            # Add translated paragraph
            translated_paragraph = translated_doc.add_paragraph()
            translated_paragraph.style = paragraph.style  # Preserve original style

            # Add the translated text with original formatting
            current_run = translated_paragraph.add_run(translated_text)
            if paragraph.runs:
                first_run = paragraph.runs[0]
                current_run.bold = first_run.bold
                current_run.italic = first_run.italic
                current_run.underline = first_run.underline
                current_run.font.name = first_run.font.name
                current_run.font.size = first_run.font.size
                current_run.font.color.rgb = first_run.font.color.rgb

            # Check for images in paragraph
            for run in paragraph.runs:
                if 'image' in run._r.xml:
                    copy_images(doc, translated_doc, translated_paragraph)

    # Save the translated document
    output_filename = f'TranslatedDocument_{uuid.uuid4()}.docx'
    output_path = os.path.join('static', output_filename)
    translated_doc.save(output_path)

    return jsonify({'message': 'Document translated successfully', 'filename': output_filename, 'path': output_path}), 200

def translate_text_logic(text, language_code):
    if not text:
        return ""
    try:
        translated = translator.translate(text, dest=language_code)
        return translated.text
    except Exception as e:
        print(f"Error during translation: {e}")
        return text

def copy_table_style(original_table, translated_table):
    for row in translated_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top="single", bottom="single", start="single", end="single")

    if original_table.style:
        translated_table.style = original_table.style

def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'start', 'bottom', 'end'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f"w:{edge}")
            element.set(qn('w:val'), edge_data)
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), 'auto')
            tcPr.append(element)

# Copy images from original to translated document
def copy_images(original_doc, translated_doc, target_element):
    """
    Copy images from the original document and insert them into the translated document.
    The `target_element` ensures the image is placed in the correct location (either table cell or paragraph).
    """
    for rel in original_doc.part.rels.values():
        if "image" in rel.target_ref:
            image_part = rel.target_part
            image_data = image_part.blob

            # Use PIL to get image dimensions
            image_stream = io.BytesIO(image_data)
            with Image.open(image_stream) as img:
                width, height = img.size

            # Convert to inches (assuming 96 DPI for Word documents)
            max_width_inch = 6  # Max width for image in inches
            max_height_inch = 8  # Max height for image in inches

            # Adjust dimensions if they exceed max size
            aspect_ratio = width / height
            if width > max_width_inch * 96:  # 96 dpi is the default in python-docx
                width = max_width_inch * 96
                height = width / aspect_ratio
            if height > max_height_inch * 96:
                height = max_height_inch * 96
                width = height * aspect_ratio

            # Insert image into the correct place in the translated document
            image_run = target_element.add_run()  # Add a run in the target paragraph or table cell
            image_run.add_picture(io.BytesIO(image_data), width=Inches(width / 96), height=Inches(height / 96))

if __name__ == '__main__':
    app.run(debug=True)
